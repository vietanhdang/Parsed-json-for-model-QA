import docx
import json
import unicodedata
import sys

TIME_LIMIT = 10

# if sys.argv.__len__() < 3:
#     print("Usage: main.py <filename.docx> <filename.json>")
#     exit(1)
def jsonEncry(self):
    word_file = docx.Document(self)
    main_dictionary = {"subjects": []}
    subject = word_file.paragraphs[0].text.split(":")[1].strip()
    no_qs = int(word_file.paragraphs[1].text.split(":")[1].strip())
    main_dictionary["subjects"].append(
        {subject: [{"number-of-question": no_qs, "time-limit": TIME_LIMIT}]}
    )

    qs_dictionary = {}
    for i, table in enumerate(word_file.tables):
        q_entry_dict = {}
        q_entry_dict["context"] = table.rows[0].cells[1].text
        a_dict = {}
        num_opts = 0
        for _ in range(1, 7):
            if table.rows[_].cells[1].text.strip() == "":
                continue
            else:
                a_dict[table.rows[_].cells[0].text[:1]] = unicodedata.normalize(
                    "NFD", table.rows[_].cells[1].text
                )
                num_opts += 1
        q_entry_dict["nums-of-options"] = str(num_opts)
        q_entry_dict["options"] = a_dict
        for _ in range(7, 11):
            q_entry_dict[
                table.rows[_].cells[0].text[0:-1].lower()
            ] = unicodedata.normalize("NFD", table.rows[_].cells[1].text.lower())

        qs_dictionary[str(i)] = q_entry_dict

    main_dictionary["subjects"][0][subject].append(qs_dictionary)
    return main_dictionary
