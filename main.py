from tkinter import *
from tkinter import filedialog
from tkinter.filedialog import askopenfilename, asksaveasfile, test
from PyPDF2 import PdfFileReader
from tkinter import ttk
import docx
from tkinter import messagebox
from docxConvertByMe import *
import json


# =================open file method======================
def openFile():
    file = askopenfilename(
        defaultextension="*.",
        filetypes=(
            ("Word 97-2003", "*.doc"),
            ("Word 2007 Until now ", "*.docx"),
            ("PDF ", "*.pdf"),
        ),
    )
    if file == "":
        file = None
    else:
        fileEntry.delete(0, END)
        fileEntry.config(fg="blue")
        fileEntry.insert(0, file)
        if fileEntry.get().split(".")[-1] == "docx":
            convertDocx()
        elif fileEntry.get().split(".")[-1] == "pdf":
            convertPdf()


def convertPdf():
    try:
        pdf = fileEntry.get()
        pdfFile = open(pdf, "rb")
        # creating a pdf reader object
        pdfReader = PdfFileReader(pdfFile)

        # creating a page object

        for i in range(0, pdfReader.numPages):
            pageObj = pdfReader.getPage(i)
            # extracting text from page
            extractedText = pageObj.extractText()
            textBox.delete(1.0, END)
            textBox.insert(INSERT, extractedText)
        # closing the pdf file object
        pdfFile.close()
    except FileNotFoundError:
        fileEntry.delete(0, END)
        fileEntry.config(fg="red")
        fileEntry.insert(0, "Please select a file first")
    except:
        pass


def convertDocx():
    try:
        word_file1 = docx.Document(fileEntry.get())
        for i in range(0, len(word_file1.paragraphs)):
            text = word_file1.paragraphs[i].text
            if text == "":
                continue
            textBox.insert(INSERT, text + "\n")

        for i, table in enumerate(word_file1.tables):
            for _ in range(0, 1):
                textBox.insert(
                    INSERT,
                    table.rows[_].cells[0].text
                    + "   "
                    + table.rows[_].cells[1].text.strip()
                    + "\n\n",
                )
            for _ in range(1, 11):
                textBox.insert(
                    INSERT,
                    table.rows[_].cells[0].text
                    + " "
                    + table.rows[_].cells[1].text.strip()
                    + "\n",
                )
    except FileNotFoundError:
        fileEntry.delete(0, END)
        fileEntry.config(fg="red")
        fileEntry.insert(0, "Please select a file first")
    except:
        pass


def export():
    modeCheck = "a"
    if Checkbutton1.get() == 1:
        modeCheck = "w"
    text = str(textBox.get(1.0, END))
    file = asksaveasfile(
        mode=modeCheck,
        filetypes=[
            ("json file", "*.json"),
        ],
    )
    if file is None:
        return
    file.write(
        str(json.dumps(jsonEncry(fileEntry.get()), indent=4, ensure_ascii=False))
    )
    file.close()
    messagebox.showinfo("Save file successfully", "Save file successfully")
    # with open(file.write(), modeCheck) as f:
    #     print(jsonEncry(fileEntry.get()), file=f)


# =================== Front End Design
root = Tk()
root.geometry("720x550")
root.config(bg="light blue")
root.title("Convert Json Tool")
root.resizable(0, 0)
try:
    root.wm_iconbitmap("pdf2.ico")
except:
    print("icon file is not available")
    pass
file = ""
# ==============App Name==============================================================>>
appName = Label(
    root,
    text="           Dang Viet Anh",
    font=("arial", 20, "bold"),
    bg="light blue",
    fg="maroon",
)
appName.place(x=200, y=5)
# ===========button to access openFile method=================================
openFileButton = Button(
    root,
    text=" Open ",
    font=("arial", 12, "bold"),
    width=5,
    bg="sky blue",
    fg="green",
    command=openFile,
)
openFileButton.place(x=30, y=35)

# Select pdf file
fileEntry = Entry(root, font=("calibri", 12), width=40, bd=4)
fileEntry.pack(ipadx=200, pady=70, padx=30)
# content
labelFile = Label(
    root,
    text="Content: (Your extracted text will apear here,you can modify that text too)",
    font=("arial", 12, "bold"),
    bg="sky blue",
    fg="green",
    bd=4,
)
labelFile.place(x=30, y=110)
# ======================= Text Box to read pdf file and modify ===================>>
textBox = Text(
    root,
    font=("calibri", 12),
    fg="light green",
    bg="black",
    width=64,
    height=15,
    bd=9,
)
textBox.place(x=30, y=145)
# ===============================Button to access export=================
Checkbutton1 = IntVar()
Overwrite = Checkbutton(
    root,
    text="Auto Overwrite",
    variable=Checkbutton1,
    onvalue=1,
    offvalue=0,
    # height=5,
    # width=20,
    bg="sky blue",
    fg="green",
)
Overwrite.place(x=30, y=460)

saveFileButton = Button(
    root,
    text=" Save ",
    font=("arial", 12, "bold"),
    width=5,
    bg="sky blue",
    fg="green",
    command=export,
)
saveFileButton.place(x=30, y=500)
# Notification save file

# fileEntry = Entry(root, font=("calibri", 12), width=40, bd=4)


# fileEntry.pack(ipadx=200, pady=400, padx=30)

# ===================halt window=============================>>
if __name__ == "__main__":
    root.mainloop()

